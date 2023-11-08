import docx
import textract
from abc import ABC
from Codigo.Model.File import File


def extract_text_from_table(table):
    table_text = ""
    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    table_text += paragraph.text + "\n"
    except Exception as e:
        return table_text

    return table_text


class DocxFile(File, ABC):
    def __init__(self, name, url_file):
        super().__init__(name, url_file)

    def get_text(self):
        try:
            #path = self.url_file.replace("/", "\\")
            #text = textract.process(path, method='docx', encoding='utf8').decode('utf8')
            #text = self.verify_footer(text)

            text = ""
            doc = docx.Document(self.url_file)
            p = 0
            t = 0

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"


            # for element in doc.element.body:
            #     if isinstance(element, docx.oxml.text.paragraph.CT_P):
            #         text += doc.paragraphs[p].text + "\n"
            #         p += 1
            #     if isinstance(element, docx.oxml.table.CT_Tbl):
            #         text += extract_text_from_table(doc.tables[t])
            #         t += 1

            # Escribe el texto en un archivo de texto
            path = self.path + self.name + ".txt"
            with open(path, "w", encoding="utf8") as f:
                f.write(text)
            return {
                'response': False,
                'message': text,

            }
        except FileNotFoundError as e:
            return {
                'response': True,
                'message': f"Error: El archivo {self.url_file} no se encontró.",
            }
        except Exception as e:
            return {
                'response': True,
                'message': f"Error: No se logró cargar el archivo.{e}         {self.url_file}",
            }

    def verify_footer(self, text):

        last_paragraph = ""
        last_paragraph_word = ""
        last_table = ""
        last_table_word = ""
        # Abre el archivo DOCX
        doc = docx.Document(self.url_file)

        # Inicializa un índice para recorrer los párrafos y las tablas hacia atrás
        index = -1
        index1 = -1
        index2 = -1

        if len(doc.paragraphs) > 0:
            # Bucle while para recorrer los párrafos hacia atrás
            while index >= -len(doc.paragraphs):
                para = doc.paragraphs[index]

                # Comprueba si el párrafo es parte del pie de página o si está vacío
                if para.text.strip() != "" and para.text != "\x0c":
                    # Agrega el texto del párrafo al principio de la cadena (hacia atrás)
                    last_paragraph = para.text
                    break
                index -= 1

            last_paragraph_word = last_paragraph.split()[-1]
            index1 = text.rindex(last_paragraph_word)

        # Inicializa un índice para recorrer las tablas hacia atrás
        index = -1
        if len(doc.tables) > 0:
            # Bucle while para recorrer las tablas hacia atrás
            while index >= -len(doc.tables):

                table = doc.tables[index]
                t = extract_text_from_table(table)

                if t.strip() != "" and t != "\x0c":
                    # Agrega el texto de la tabla al principio de la cadena (hacia atrás)
                    last_table = t
                    break
                index -= 1

            if last_table != "":
                last_table_word = last_table.split()[-1]
                index2 = text.rindex(last_table_word)

        if index1 > index2:
            text = text[:index1] + last_paragraph_word
        elif index2 > index1:
            text = text[:index2] + last_table_word

        return text
